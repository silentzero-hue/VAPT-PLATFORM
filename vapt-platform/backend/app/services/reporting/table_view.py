"""Per-severity table view renderer.

This is the exact feature the spec was hinting at when it said:
"the same vulnerability appearing across multiple host groups,
split into separate severity tables in the source report, must
dedup into one vulnerability record with multiple linked findings".

The legacy tool produced tables like:

  ┌─ CRITICAL FINDINGS ──────────────────────────────────────┐
  │ CVE          │ Title                │ Hosts │ Port │ CVSS │
  ├──────────────┼──────────────────────┼───────┼──────┼──────┤
  │ CVE-2024-001 │ Log4Shell on Tomcat  │ 12    │ 443  │ 10.0 │
  │ CVE-2024-002 │ OpenSSL EOL          │ 8     │ 443  │ 9.8  │
  └──────────────┴──────────────────────┴───────┴──────┴──────┘

We render these as docx tables, one per severity band, with the
deduplicated Vulnerability (one row per unique CVE) and the count
of hosts it manifested on.
"""

from __future__ import annotations

import io
import uuid
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.engagement import Engagement
from app.models.finding import Finding
from app.models.vulnerability import Severity, Vulnerability


SEV_BAND_ORDER = [
    (Severity.CRITICAL, "Critical"),
    (Severity.HIGH, "High"),
    (Severity.MEDIUM, "Medium"),
    (Severity.LOW, "Low"),
    (Severity.INFO, "Informational"),
]

SEV_COLOR_RGB = {
    Severity.CRITICAL: RGBColor(0xFF, 0x4D, 0x6D),
    Severity.HIGH:     RGBColor(0xFF, 0x8C, 0x42),
    Severity.MEDIUM:   RGBColor(0xFF, 0xD1, 0x66),
    Severity.LOW:      RGBColor(0x3D, 0xDC, 0x97),
    Severity.INFO:     RGBColor(0x7A, 0xA1, 0xFF),
}


async def build_table_view(
    db: AsyncSession, engagement_id: uuid.UUID
) -> dict[str, Any]:
    """Return the table data structured by severity band. Each row is:
        {
          "vuln_id": str,
          "cve_id": str | None,
          "title": str,
          "cvss": float | None,
          "host_count": int,
          "host_list": [str, ...],
          "port_list": [int, ...],
          "sample_asset": str,
        }
    """
    e = await db.get(Engagement, engagement_id)
    if not e:
        return {}
    findings = (await db.execute(
        select(Finding).where(Finding.engagement_id == engagement_id)
    )).scalars().all()
    # group by vuln
    by_vuln: dict[uuid.UUID, list[Finding]] = {}
    for f in findings:
        by_vuln.setdefault(f.vulnerability_id, []).append(f)

    by_band: dict[str, list[dict]] = {sev_enum.value: [] for sev_enum, _ in SEV_BAND_ORDER}
    for vid, fs in by_vuln.items():
        v = await db.get(Vulnerability, vid)
        if not v:
            continue
        from app.models.asset import Asset
        hosts: set[str] = set()
        ports: set[int] = set()
        sample_asset = ""
        for f in fs:
            a = await db.get(Asset, f.asset_id)
            if a:
                hosts.add(a.value)
                sample_asset = sample_asset or a.value
            if f.port is not None:
                ports.add(f.port)
        row = {
            "vuln_id": str(v.id),
            "cve_id": v.cve_id,
            "title": v.title,
            "cvss_score": v.cvss_score,
            "host_count": len(hosts),
            "hosts": sorted(hosts)[:20],
            "ports": sorted(p for p in ports if p is not None),
            "sample_asset": sample_asset,
            "occurrence_count": v.occurrence_count,
        }
        if v.severity and v.severity.value in by_band:
            by_band[v.severity.value].append(row)
    return {
        "engagement": {"code": e.code, "name": e.name, "client": e.client},
        "by_severity": by_band,
        "totals": {label: len(rows) for label, rows in by_band.items()},
    }


def render_table_view_docx(data: dict[str, Any]) -> bytes:
    """Render the table view to a docx file. One table per severity band."""
    if not data:
        doc = Document()
        doc.add_paragraph("No data")
        buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

    doc = Document()
    e = data.get("engagement", {})
    doc.add_heading(f"{e.get('name', 'Engagement')} — Findings by Severity", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Client: {e.get('client', '')}    Code: {e.get('code', '')}").italic = True

    for sev_enum, label in SEV_BAND_ORDER:
        rows = data["by_severity"].get(sev_enum.value, [])
        if not rows:
            continue
        doc.add_heading(f"{label} ({len(rows)})", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["CVE / ID", "Title", "CVSS", "Hosts", "Ports", "Sample asset"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for r in rows:
            row = table.add_row().cells
            row[0].text = r.get("cve_id") or r.get("vuln_id", "")[:8]
            row[1].text = (r.get("title") or "")[:120]
            row[2].text = f"{r.get('cvss_score'):.1f}" if r.get("cvss_score") else "—"
            row[3].text = str(r.get("host_count", 0))
            row[4].text = ", ".join(str(p) for p in r.get("ports", [])[:6])
            row[5].text = r.get("sample_asset", "—")
            for c in row:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
        # color the heading
        for p in doc.paragraphs:
            if p.text.startswith(label):
                for run in p.runs:
                    run.font.color.rgb = SEV_COLOR_RGB[sev_enum]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_table_view_html(data: dict[str, Any]) -> str:
    """Render the table view as a self-contained HTML page (for the
    web UI). No external CSS — we inline the styles so the browser
    can print-to-PDF directly."""
    if not data:
        return "<p>No data</p>"
    e = data.get("engagement", {})
    out = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        "<title>", e.get("name", "Findings"), "</title>",
        "<style>",
        "body{font-family:Inter,system-ui,sans-serif;max-width:1100px;margin:24px auto;color:#0b1020;background:#fff;padding:0 16px}",
        "h1{font-size:22px;margin-bottom:4px}",
        "h2{margin-top:32px;padding:6px 10px;border-radius:6px;color:#fff}",
        ".critical{background:#ff4d6d} .high{background:#ff8c42} .medium{background:#ffd166} .low{background:#3ddc97} .info{background:#7aa1ff}",
        "table{border-collapse:collapse;width:100%;margin:8px 0 24px}",
        "th,td{border:1px solid #dbe0ee;padding:6px 8px;text-align:left;font-size:13px}",
        "th{background:#eef1f8}",
        ".cve{font-family:ui-monospace,monospace;font-size:12px}",
        ".muted{color:#666}",
        "</style></head><body>",
        f"<h1>{e.get('name', '')}</h1>",
        f"<p class='muted'>Client: {e.get('client', '')} &middot; Code: {e.get('code', '')}</p>",
    ]
    for sev_enum, label in SEV_BAND_ORDER:
        rows = data["by_severity"].get(sev_enum.value, [])
        if not rows:
            continue
        cls = label.lower() if label != "Informational" else "info"
        out.append(f"<h2 class='{cls}'>{label} ({len(rows)})</h2>")
        out.append("<table><thead><tr>")
        for h in ["CVE / ID", "Title", "CVSS", "Hosts", "Ports", "Sample asset"]:
            out.append(f"<th>{h}</th>")
        out.append("</tr></thead><tbody>")
        for r in rows:
            cve = r.get("cve_id") or (r.get("vuln_id") or "")[:8]
            ports = ", ".join(str(p) for p in r.get("ports", [])[:6])
            cvss = f"{r.get('cvss_score'):.1f}" if r.get("cvss_score") else "—"
            out.append(
                f"<tr><td class='cve'>{cve}</td>"
                f"<td>{(r.get('title') or '')[:120]}</td>"
                f"<td>{cvss}</td><td>{r.get('host_count', 0)}</td>"
                f"<td>{ports}</td><td>{r.get('sample_asset', '—')}</td></tr>"
            )
        out.append("</tbody></table>")
    out.append("</body></html>")
    return "".join(out)
