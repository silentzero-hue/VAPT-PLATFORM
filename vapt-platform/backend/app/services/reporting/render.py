"""DOCX report rendering service.

Renders VAPT reports using the Technovage Solution DMC template. The
template lives at ``app/services/reporting/templates/dmc_vapt_report.docx``
and is loaded with python-docx (the template has no Jinja placeholders,
so docxtpl is not used here). If the template file is missing we fall
back to building the report from scratch with python-docx defaults.
"""

from __future__ import annotations

import io
import json
import os
import re
import structlog
import uuid
from datetime import datetime, timezone
from typing import Any

import mammoth
from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from lxml import html as lxml_html
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.asset import Asset
from app.models.engagement import Engagement
from app.models.finding import Finding
from app.models.vulnerability import Vulnerability
from app.services.reporting.suggestions import URGENCY_LABEL as ACTION_URGENCY

SEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

# Where the bundled DMC template lives inside the source tree.
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
TEMPLATE_FILENAME = "dmc_vapt_report.docx"

# Scanner-output patterns to strip from finding descriptions before they
# reach the rendered docx. These tags come straight from Nessus plugin XML
# (e.g. ``<text>[description]</text>`` wraps the prose). We keep CVE
# references, ``Tenable Plugin ID:``, ``Tenable Reference:``, and CWE
# references — those are useful to a human reader. We strip bare
# ``Plugin ID:`` (without the ``Tenable`` prefix) and the standalone
# ``Plugin - <digits>`` lines, which are scan-tool internal refs.
_SCANNER_OUTPUT_PATTERNS: list[re.Pattern[str]] = [
    re.compile(r"\[/description\]\s*", re.I),
    re.compile(r"\[description\]\s*", re.I),
    re.compile(r"\[/synopsis\]\s*", re.I),
    re.compile(r"\[synopsis\]\s*", re.I),
    re.compile(r"\[/plugin_output\]\s*", re.I),
    re.compile(r"\[plugin_output\]\s*", re.I),
    re.compile(r"\[/solution\]\s*", re.I),
    re.compile(r"\[solution\]\s*", re.I),
    re.compile(r"\[/see_also\]\s*", re.I),
    re.compile(r"\[see_also\]\s*", re.I),
    re.compile(r"\[/risk_factor\]\s*", re.I),
    re.compile(r"\[risk_factor\]\s*", re.I),
    re.compile(r"\[/cvss_base_vector\]\s*", re.I),
    re.compile(r"\[cvss_base_vector\]\s*", re.I),
    re.compile(r"\[/cvss3_base_vector\]\s*", re.I),
    re.compile(r"\[cvss3_base_vector\]\s*", re.I),
    re.compile(r"^\s*Plugin\s*-\s*\d+\s*$", re.M | re.I),
    re.compile(r"^\s*Plugin\s*ID\s*:\s*\d+\s*$", re.M | re.I),
    re.compile(r"\bPlugin\s*-\s*\d+\b", re.I),
    re.compile(r"^\s*Description\s*$\s*", re.M | re.I),
    re.compile(r"^\s*Plugin Output\s*$\s*", re.M | re.I),
]


def _clean_scanner_output(text: str) -> str:
    """Remove scanner-output markers and boilerplate from a description.

    The underlying Nessus parser stores the raw plugin XML body as the
    vulnerability's ``description`` field. That body is full of marker
    tags (``[/description]``, ``[synopsis]`` …) and ``Plugin - 12345``
    references that should never reach a client-facing report.

    What we keep (intentionally):
      * CVE references (``CVE-2024-1234``)
      * ``Tenable Plugin ID:``, ``Tenable Plugin:``, ``Tenable Reference:`` —
        these are useful evidence for the analyst.
      * ``CWE:`` references.
      * The actual prose explaining the vulnerability.

    What we strip:
      * All ``[xxx]`` / ``[/xxx]`` wrapper tags.
      * Standalone ``Plugin - 12345`` lines.
      * Standalone ``Plugin ID: 12345`` lines (the ``Tenable Plugin ID:``
        form is preserved).
      * Multiple blank lines (collapsed to a single blank line).
    """
    if not text:
        return ""
    cleaned = text
    for pat in _SCANNER_OUTPUT_PATTERNS:
        cleaned = pat.sub("", cleaned)
    # Collapse runs of 3+ newlines into a single blank line
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    # Collapse runs of 2+ spaces (artifacts from tag removal) inside lines
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    # Strip leading/trailing whitespace
    cleaned = cleaned.strip()
    return cleaned


def _software_key(title: str) -> str:
    """Extract a 'software key' from a finding title for merging purposes.

    The point is to collapse findings that refer to the same software
    even when they have different Nessus plugin IDs / CVE IDs — e.g.
    ``OpenSSH < 7.8``, ``OpenSSH < 8.0``, and
    ``OpenSSH < 9.6 Multiple Vulnerabilities (CVE-2023-48795)`` should
    all merge into a single finding.

    Examples:
        "OpenSSH < 7.8"                                  -> "openssh"
        "OpenSSH < 8.0"                                  -> "openssh"
        "OpenSSH < 9.6 Multiple Vulnerabilities (...)"   -> "openssh"
        "Apache ActiveMQ RCE (CVE-2023-46604)"            -> "apache activemq"
        "RHEL 8 : bzip2 (RHSA-2025:0733)"                 -> "rhel 8"
        "PHP < 8.1 Multiple Vulnerabilities"              -> "php"
    """
    if not title:
        return ""
    t = title.lower().strip()
    # First word must be a letter (skip "1. ", "[", etc.)
    m = re.match(r"^([a-z][a-z0-9]+(?:\s+[a-z0-9]+){0,2})", t)
    if m:
        return m.group(1).strip()
    return t


def _merge_similar_findings(groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Merge groups that share the same software, even if they have
    different ``vulnerability_id``s.

    The DMC report collapses anything that mentions the same piece of
    software into one finding — e.g. all OpenSSH CVEs in one row. We
    replicate that here. The merge keeps:
      * the highest-severity group as the "primary" (its title, port,
        impact, recommendation, etc.),
      * the longest descriptive title (most-specific wins),
      * the union of CVEs and affected hosts across all merged groups.

    Findings whose title doesn't yield a software key are kept as-is
    (no merge), so a poorly-titled finding still surfaces.
    """
    by_software: dict[str, list[dict[str, Any]]] = {}
    order: list[str] = []
    for grp in groups:
        sw = _software_key(grp.get("issues_title") or grp.get("title") or "")
        if not sw:
            # No key — give it a unique placeholder so it doesn't merge
            sw = f"__no_key_{id(grp)}"
        if sw not in by_software:
            by_software[sw] = []
            order.append(sw)
        by_software[sw].append(grp)

    merged: list[dict[str, Any]] = []
    for sw in order:
        grps = by_software[sw]
        if len(grps) == 1:
            merged.append(grps[0])
            continue
        # Collect combined CVEs and hosts
        combined_cves: set[str] = set()
        for g in grps:
            for cve in (g.get("cve_ids") or []):
                if cve:
                    combined_cves.add(str(cve))
            single = g.get("cve_id")
            if single and single not in combined_cves:
                combined_cves.add(str(single))
        combined_hosts: set[str] = set()
        for g in grps:
            for h in g.get("affected_assets") or []:
                if h:
                    combined_hosts.add(str(h))
        # Worst severity wins
        worst = min(grps, key=lambda g: SEV_ORDER.get((g.get("severity") or "info").lower(), 9))
        # Most descriptive title wins
        best_title = max(
            (g.get("issues_title") or g.get("title") or "") for g in grps
        )
        merged.append({
            **worst,
            "cve_ids": sorted(combined_cves),
            "cve_id": ", ".join(sorted(combined_cves)) or worst.get("cve_id", ""),
            "affected_assets": sorted(combined_hosts),
            "issues_title": best_title or worst.get("issues_title") or "",
            "merged_count": len(grps),
        })
    return merged


def _build_stub_vuln(title: str, cve_id: str | None, severity: str) -> Any:
    """Build a minimal stub object compatible with the ``suggest()`` API.

    The ``app.services.reporting.suggestions`` module expects a
    ``Vulnerability``-shaped object with ``.title``, ``.cve_id``,
    ``.severity.value``. We don't have a real Vulnerability here (this
    runs after grouping by software), so we synthesize a SimpleNamespace
    with the same surface.
    """
    from types import SimpleNamespace

    return SimpleNamespace(
        title=title or "(untitled)",
        cve_id=cve_id,
        severity=SimpleNamespace(value=(severity or "info").lower()),
    )


def _resolve_template_path() -> str:
    """Return absolute path to the DMC template, or empty string if missing."""
    raw = settings.report_template_path or ""
    if not raw:
        return ""
    # Absolute path from settings
    if os.path.isabs(raw) and os.path.exists(raw):
        return raw
    # Try relative to CWD (e.g. inside the container the workdir is /app)
    cwd_candidate = os.path.abspath(raw)
    if os.path.exists(cwd_candidate):
        return cwd_candidate
    # Fall back to bundled template (source-tree path)
    bundled = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
    if os.path.exists(bundled):
        return bundled
    return ""


async def build_report_context(db: AsyncSession, engagement_id: uuid.UUID) -> dict[str, Any]:
    """Aggregate everything the template needs."""
    e = await db.get(Engagement, engagement_id)
    findings = (
        (await db.execute(select(Finding).where(Finding.engagement_id == engagement_id)))
        .scalars()
        .all()
    )
    vuln_ids = {f.vulnerability_id for f in findings}
    vulns: dict[uuid.UUID, Vulnerability] = {}
    for vid in vuln_ids:
        v = await db.get(Vulnerability, vid)
        if v:
            vulns[vid] = v

    # severity histogram
    sev_count: dict[str, int] = {}
    for f in findings:
        v_for_f = vulns.get(f.vulnerability_id)
        sev = f.severity_override or (v_for_f.severity.value if v_for_f else "info")
        sev_count[sev] = sev_count.get(sev, 0) + 1

    # group by vuln (deduplicate narrative per vuln)
    by_vuln: dict[uuid.UUID, list[Finding]] = {}
    for f in findings:
        by_vuln.setdefault(f.vulnerability_id, []).append(f)

    # build top-level finding list, one row per (vuln, asset) but with shared narrative
    rows: list[dict[str, Any]] = []
    # for asset values
    asset_cache: dict[uuid.UUID, Asset] = {}
    for f in findings:
        if f.asset_id not in asset_cache:
            a = await db.get(Asset, f.asset_id)
            if a:
                asset_cache[f.asset_id] = a
        a = asset_cache.get(f.asset_id)
        v = vulns.get(f.vulnerability_id)
        if not v:
            continue
        rows.append(
            {
                "finding_id": str(f.id),
                "vuln_id": str(v.id),
                "title": v.title,
                "cve_id": v.cve_id,
                "cwe_id": v.cwe_id,
                "severity": (f.severity_override or v.severity.value),
                "cvss_score": f.cvss_score_override or v.cvss_score,
                "asset_value": a.value if a else "?",
                "asset_type": a.type.value if a else "host",
                "port": f.port,
                "protocol": f.protocol,
                "status": f.status.value,
                "evidence": (f.evidence_ref or "")[:400],
                "description": v.description,
                "impact": v.ai_draft_impact or "",
                "recommendation": v.ai_draft_recommendation or "",
                "linked_asset_count": len(by_vuln.get(v.id, [])),
                "source_plugin": v.source_plugin,
                "source_plugin_id": v.source_plugin_id,
                "references": list(v.references or []),
                "issues_text": _format_issues_cell(v, f),
            }
        )
    rows.sort(key=lambda r: (SEV_ORDER.get(r["severity"], 9), -r["linked_asset_count"]))

    # Group findings by (vuln, port) for the Detailed Findings table. Each
    # group produces one row whose "Affected IP/URL" cell lists every
    # asset affected by that vuln.
    vuln_groups: dict[tuple[str, str], dict[str, Any]] = {}
    for f in findings:
        a = asset_cache.get(f.asset_id) or await db.get(Asset, f.asset_id)
        if a:
            asset_cache[f.asset_id] = a
        v = vulns.get(f.vulnerability_id)
        if not v:
            continue
        port_key = "" if f.port is None else str(f.port)
        key = (str(v.id), port_key)
        grp = vuln_groups.setdefault(
            key,
            {
                "vuln_id": str(v.id),
                "title": v.title,
                "cve_id": v.cve_id,
                "severity": (f.severity_override or v.severity.value),
                "port": f.port,
                "protocol": f.protocol,
                "impact": v.ai_draft_impact or "",
                "recommendation": v.ai_draft_recommendation or "",
                "assets": [],
                "description": v.description,
                "source_plugin": v.source_plugin,
                "source_plugin_id": v.source_plugin_id,
                "references": list(v.references or []),
                "issues_text": _format_issues_cell(v, f),
            },
        )
        grp["assets"].append(a.value if a else "?")

    detailed_groups: list[dict[str, Any]] = []
    for grp in vuln_groups.values():
        detailed_groups.append(
            {
                "vuln_id": grp["vuln_id"],
                "severity": grp["severity"],
                "affected_assets": grp["assets"],
                "port": grp["port"],
                "protocol": grp["protocol"],
                "issues_title": grp["title"],
                "cve_id": grp["cve_id"],
                "impact": grp["impact"],
                "recommendation": grp["recommendation"],
                "action_urgency": ACTION_URGENCY.get(grp["severity"], "Normal"),
                "description": grp["description"],
                "source_plugin": grp["source_plugin"],
                "source_plugin_id": grp["source_plugin_id"],
                "references": grp["references"],
                "issues_text": grp["issues_text"],
            }
        )

    # Second pass: merge findings that share the same software, even
    # when they have different vulnerability_ids. The DMC report's
    # style is "OpenSSH < 7.8" + "OpenSSH < 8.0" + "OpenSSH Terrapin"
    # all in one row, with combined CVEs and affected hosts.
    detailed_groups = _merge_similar_findings(detailed_groups)

    # Fallback to the category-specific suggestions engine for impact
    # and recommendation when the AI drafts are empty or too short to
    # be useful. This avoids the boilerplate "See Nessus plugin output
    # for technical details." and "Apply vendor patch per Nessus
    # recommendation." text the old renderer was emitting.
    from app.services.reporting.suggestions import suggest as _suggest
    for grp in detailed_groups:
        title = grp.get("issues_title") or grp.get("title") or "(untitled)"
        cve_id = grp.get("cve_id") or None
        severity = (grp.get("severity") or "info").lower()
        if not (grp.get("impact") or "").strip():
            stub = _build_stub_vuln(title=title, cve_id=cve_id, severity=severity)
            sug = _suggest(stub, severity_override=severity)
            grp["impact"] = sug["impact"]
        if not (grp.get("recommendation") or "").strip():
            stub = _build_stub_vuln(title=title, cve_id=cve_id, severity=severity)
            sug = _suggest(stub, severity_override=severity)
            grp["recommendation"] = sug["recommendation"]

    detailed_groups.sort(key=lambda r: (SEV_ORDER.get(r["severity"], 9), r["issues_title"]))

    return {
        "engagement": {
            "code": e.code if e else "",
            "name": e.name if e else "",
            "client": e.client if e else "",
            "type": e.type.value if e else "",
            "status": e.status.value if e else "",
            "methodology": e.methodology if e else None,
            "start_date": str(e.start_date) if e and e.start_date else "",
            "end_date": str(e.end_date) if e and e.end_date else "",
            "report_due_date": str(e.report_due_date) if e and e.report_due_date else "",
        },
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "summary": {
            "total_findings": len(findings),
            "by_severity": sev_count,
            "unique_vulnerabilities": len(vulns),
            "assets_affected": len({f.asset_id for f in findings}),
        },
        "findings": rows,
        "detailed_findings": detailed_groups,
    }


def _apply_draft(ctx: dict[str, Any], draft: dict[str, Any]) -> None:
    """Mutate `ctx` in place to apply the analyst's draft overrides.

    Mapping rules:
    - draft["overall_rating"]  ->  ctx["summary"]["overall_rating"]
    - draft["exec_summary"]    ->  ctx["exec_summary"]
    - draft["finding_overrides"][finding_id] (str-keyed UUID):
        - severity_override  ->  finding row + matching detailed row
        - impact, recommendation, note  ->  finding row only
        (the detailed group uses the finding's overridden impact/rec
         so the template's per-row cells pick them up)
    """
    if not draft:
        return

    overall = draft.get("overall_rating")
    if isinstance(overall, str) and overall:
        ctx.setdefault("summary", {})["overall_rating"] = overall

    es = draft.get("exec_summary")
    if isinstance(es, str) and es:
        ctx["exec_summary"] = es

    overrides = draft.get("finding_overrides") or {}
    if not isinstance(overrides, dict) or not overrides:
        return

    # build a quick map for detailed_findings: (vuln_id, port_str) -> group
    detail_index: dict[tuple[str, str], dict[str, Any]] = {}
    for grp in ctx.get("detailed_findings", []):
        vid = grp.get("vuln_id") or ""
        port = grp.get("port")
        port_key = "" if port is None else str(port)
        detail_index[(vid, port_key)] = grp

    for row in ctx.get("findings", []):
        fid = str(row.get("finding_id") or "")
        ov = overrides.get(fid)
        if not isinstance(ov, dict):
            continue
        sev = ov.get("severity_override")
        if isinstance(sev, str) and sev:
            row["severity"] = sev
        impact = ov.get("impact")
        if isinstance(impact, str):
            row["impact"] = impact
        rec = ov.get("recommendation")
        if isinstance(rec, str):
            row["recommendation"] = rec
        note = ov.get("note")
        if isinstance(note, str):
            row["note"] = note
        # Propagate impact/rec into the matching detailed_findings group
        # so the template's per-row cells reflect the same override.
        vid = str(row.get("vuln_id") or "")
        port = row.get("port")
        port_key = "" if port is None else str(port)
        grp = detail_index.get((vid, port_key))
        if grp is not None:
            if isinstance(impact, str):
                grp["impact"] = impact
            if isinstance(rec, str):
                grp["recommendation"] = rec
            if isinstance(sev, str) and sev:
                grp["severity"] = sev
                grp["action_urgency"] = ACTION_URGENCY.get(sev.lower(), "Normal")


# ---------------------------------------------------------------------------
# helpers for filling the DMC template
# ---------------------------------------------------------------------------


def _set_cell_text(cell: _Cell, text: str) -> None:
    """Replace the contents of a docx cell with a single line of text,
    preserving the cell's existing paragraph/character formatting."""
    text = text or ""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(text)
        return
    first = paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    # Drop any extra paragraphs that came with the template.
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _set_cell_multiline(cell: _Cell, text: str) -> None:
    """Replace cell contents with text split on ``\\n`` across paragraphs.

    Defensive: any paragraph that contains a ``<w:drawing>`` or ``<w:pict>``
    (inline image or VML pict) is PRESERVED — the cell is rewritten around
    it. The original DMC cover-page right cell does not have inline images,
    but the renderer must never destroy one if a future template variant
    embeds the Technovage logo there.
    """
    text = text or ""
    lines = text.split("\n") if text else [""]
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph(lines[0])
        paragraphs = list(cell.paragraphs)

    # Identify the paragraphs to KEEP: the first text-only one, plus any
    # paragraph that contains an inline image (drawing or pict). Image-bearing
    # paragraphs are removed from the list of "text slots" but their XML
    # elements are left in the cell.
    drawing_qn = qn("w:drawing")
    pict_qn = qn("w:pict")

    def _has_image(p_el: Any) -> bool:
        return bool(p_el.findall(".//" + drawing_qn) or p_el.findall(".//" + pict_qn))

    text_slots: list[Any] = []
    keep_indices: list[int] = []
    for i, p in enumerate(paragraphs):
        if _has_image(p._element):
            keep_indices.append(i)
        else:
            text_slots.append(p)
            if not keep_indices and i == 0:
                keep_indices.append(0)

    if not text_slots:
        # The cell is entirely image-bearing (very unusual). Bail — don't
        # touch the cell, the analyst will need to look at the template.
        return

    first = text_slots[0]
    if first.runs:
        first.runs[0].text = lines[0]
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(lines[0])

    # Remove all text-slot paragraphs except the first. The image-bearing
    # paragraphs are never in this list, so they survive intact.
    for p in text_slots[1:]:
        p._element.getparent().remove(p._element)

    # Add the remaining lines as new paragraphs. We append them after the
    # LAST existing element in the cell, so they end up at the end of the
    # cell — image paragraphs, if any, stay in their original position.
    for line in lines[1:]:
        cell.add_paragraph(line)


def _format_issues_cell(vuln: Vulnerability, finding: Finding) -> str:
    """Format the Issues column for the Detailed Findings table.

    Mirrors the structure used in real DMC reports:
      <title> (CVE-XXXX)
      <empty line>
      <plugin description>
      <empty line>
      Tenable Plugin ID: <id>
      Tenable Plugin: <source_plugin>
      Tenable Reference: https://www.tenable.com/plugins/index.php?id=<id>
      CWE: CWE-XXX

    The description body is run through ``_clean_scanner_output`` so that
    Nessus XML markers (``[/description]``, ``[synopsis]``, etc.) and
    bare ``Plugin - 12345`` lines never reach the client-facing report.
    """
    parts: list[str] = []
    title = vuln.title or "(no title)"
    if vuln.cve_id:
        parts.append(f"{title} ({vuln.cve_id})")
    else:
        parts.append(title)

    # Plugin description (the "The remote host contains..." text).
    # Strip scanner-output markers before rendering.
    if vuln.description:
        cleaned = _clean_scanner_output(vuln.description)
        if cleaned:
            parts.append("")
            parts.append(cleaned)

    # References block
    refs: list[str] = []
    if vuln.source_plugin_id and vuln.source_plugin == "nessus":
        refs.append(f"Tenable Plugin ID: {vuln.source_plugin_id}")
        refs.append(f"Tenable Plugin: {vuln.source_plugin}")
        if vuln.source_plugin_id.isdigit():
            refs.append(
                f"Tenable Reference: "
                f"https://www.tenable.com/plugins/index.php?id={vuln.source_plugin_id}"
            )
    elif vuln.source_plugin:
        refs.append(f"Source plugin: {vuln.source_plugin}")
        if vuln.source_plugin_id:
            refs.append(f"Plugin ID: {vuln.source_plugin_id}")

    if vuln.cwe_id:
        refs.append(f"CWE: {vuln.cwe_id}")

    if vuln.references:
        # Avoid duplicates with the Tenable Reference line
        seen = {r.lower() for r in refs}
        for r in vuln.references:
            if r and r.lower() not in seen:
                refs.append(r)
                seen.add(r.lower())

    if refs:
        parts.append("")
        parts.extend(refs)

    return "\n".join(parts)


def _format_date(iso: str) -> str:
    """Render an ISO date string as '2nd October, 2025' style."""
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
    except ValueError:
        return iso
    day = dt.day
    suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suffix} {dt.strftime('%B, %Y')}"


def _type_label(eng_type: str) -> str:
    """Map engagement.type enum value to a human label."""
    overrides = {
        "webapp": "Web Application Penetration Testing",
        "network": "Internal Network Penetration Testing",
        "external_network": "External Network Penetration Testing",
        "wireless": "Wireless Network Penetration Testing",
        "mobile": "Mobile Application Penetration Testing",
        "social_engineering": "Social Engineering Assessment",
        "physical": "Physical Security Assessment",
        "redteam": "Red Team Assessment",
    }
    return overrides.get(eng_type.lower(), eng_type.replace("_", " ").title())


def _overall_rating(by_severity: dict[str, int]) -> str:
    """Compute overall security rating from a {severity: count} histogram."""
    if by_severity.get("critical", 0) > 0:
        return "Critical"
    if by_severity.get("high", 0) > 0:
        return "High"
    if by_severity.get("medium", 0) > 0:
        return "Medium"
    if by_severity.get("low", 0) > 0:
        return "Low"
    return "Informational"


def _looks_like_for_field(cell: _Cell) -> bool:
    """Return True if the cell holds the DMC cover-page "For / Type / From / Date"
    label pattern. Detection is loose: the cell must mention at least two of the
    four labels (case-insensitive) and the token ``:`` (a colon). This prevents
    us from ever mistaking a regular text cell for the cover-page label cell.
    """
    if cell is None:
        return False
    text = cell.text.lower()
    if ":" not in text:
        return False
    labels = ["for", "type", "from", "date"]
    hits = sum(1 for lbl in labels if lbl in text)
    return hits >= 2


def _set_left_label_value(cell: _Cell, label: str, value: str) -> None:
    """Find the first paragraph in `cell` whose first non-whitespace token
    (case-insensitive, up to the colon) matches `label` and rewrite it in
    place to ``"<label>: <value>"``. Preserves the leading whitespace of
    the original paragraph (the DMC template uses 4-char padding before
    each label, which doubles as the right-cell monospace column gutter).
    If no matching paragraph exists, the value is appended as a new
    paragraph at the end of the cell.

    Whitespace inside the matched paragraph's runs is preserved as much as
    possible: we rewrite only the *text* of the paragraph, leaving the run
    structure intact so character formatting (the cover-page's monospace
    label style) is not lost.
    """
    if cell is None or not label:
        return
    # Normalize: strip the trailing colon from `label` so we can match
    # against the first token of the paragraph (which itself may have a
    # colon attached, e.g. ``"For    :"``).
    label_word = label.rstrip(":").strip().lower()
    paragraphs = cell.paragraphs
    for p in paragraphs:
        ptext = p.text or ""
        stripped = ptext.strip()
        if not stripped:
            continue
        # Take the first whitespace-delimited token, lowercase it, and
        # strip a trailing colon if present. This handles both
        # ``"For    :"`` (DMC) and ``"For:"`` (any custom template) the
        # same way.
        first_token = stripped.split()[0].rstrip(":").lower()
        if first_token == label_word:
            leading = ptext[: len(ptext) - len(ptext.lstrip())]
            new_text = f"{leading}{label_word.title()}: {value}".rstrip()
            if p.runs:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.add_run(new_text)
            return
    # No matching paragraph — append a new line at the end of the cell.
    cell.add_paragraph(f"{label_word.title()}: {value}".rstrip())


def _fill_title_page(doc: DocumentObject, ctx: dict[str, Any]) -> None:
    """Fill table 0 (title page). The DMC template uses a 2-column cover table:
    the LEFT cell holds the ``For : / Type : / From : / Date :`` label columns
    (with the values appended after the colon), and the RIGHT cell holds the
    same values formatted as a multi-line stack (client name, type label,
    overall-rating subtitle, signing company, date)."""
    if not doc.tables:
        return
    t = doc.tables[0]
    if len(t.rows) < 1 or len(t.rows[0].cells) < 2:
        return
    eng = ctx["engagement"]
    overall_raw = ctx["summary"].get("overall_rating")
    if isinstance(overall_raw, str) and overall_raw:
        overall = overall_raw.title()
    else:
        overall = _overall_rating(ctx["summary"]["by_severity"])
    title_date = _format_date(eng.get("start_date") or "")
    # LEFT cell: append values to the "For/Type/From/Date" label paragraphs.
    # ``test_for`` is the engagement's stated test purpose; fall back to a
    # sensible default so the cell is never empty.
    left_cell = t.rows[0].cells[0]
    if _looks_like_for_field(left_cell):
        _set_left_label_value(left_cell, "For:", eng.get("name") or "Penetration Testing")
        _set_left_label_value(left_cell, "Type:", _type_label(eng.get("type") or ""))
        _set_left_label_value(
            left_cell,
            "From:",
            (settings.report_company_name or "").strip() or "the testing organization",
        )
        _set_left_label_value(left_cell, "Date:", title_date)
    # RIGHT cell: same data, formatted as a multi-line stack (the DMC style).
    lines = [
        eng.get("client") or "",
        _type_label(eng.get("type") or ""),
        f"[{overall} Findings]",
        (settings.report_company_name or "").strip() or "the testing organization",
        title_date,
    ]
    _set_cell_multiline(t.rows[0].cells[1], "\n".join(lines))


def _fill_scope_table(doc: DocumentObject, ctx: dict[str, Any]) -> None:
    """Fill table 1 (scope). Keep the header row, replace the single
    sample data row with one row per asset."""
    if len(doc.tables) < 2:
        return
    t = doc.tables[1]
    if len(t.rows) < 1:
        return
    # Determine how the data row groups IPs (the template groups many IPs
    # in a single cell, but the report schema is one row per asset).
    assets: list[str] = []
    for row in ctx.get("findings", []):
        v = row.get("asset_value")
        if v and v not in assets:
            assets.append(v)
    if not assets:
        assets = [ctx["engagement"].get("name", "")]
    test_location = "TS Office"  # template default; could be made configurable
    eng_type = ctx["engagement"].get("type", "")
    test_type = "Black Box" if "external" not in eng_type.lower() else "External"
    start_date = _format_date(ctx["engagement"].get("start_date") or "")
    # Remove all rows after the header, then add one row per asset
    for r in list(t.rows)[1:]:
        r._element.getparent().remove(r._element)
    for asset in assets:
        new_row = t.add_row()
        _set_cell_text(new_row.cells[0], asset)
        _set_cell_text(new_row.cells[1], test_location)
        _set_cell_text(new_row.cells[2], test_type)
        _set_cell_text(new_row.cells[3], start_date)


def _build_host_summary(ctx: dict[str, Any]) -> list[dict[str, Any]]:
    """Aggregate findings by (hostname, IP) and produce one row per host."""
    # asset_cache: value -> (hostname_or_ip, ip_only)
    asset_cache: dict[str, tuple[str, str]] = {}
    findings = ctx.get("findings", [])
    for row in findings:
        v = row.get("asset_value", "?")
        if v not in asset_cache:
            asset_cache[v] = (v, v)
    # counts[asset_value][severity] = n
    counts: dict[str, dict[str, int]] = {}
    for row in findings:
        v = row.get("asset_value", "?")
        sev = row.get("severity", "info").lower()
        slot = counts.setdefault(v, {})
        slot[sev] = slot.get(sev, 0) + 1
    out: list[dict[str, Any]] = []
    for i, (asset, _) in enumerate(asset_cache.items(), 1):
        c = counts.get(asset, {})
        # The "Hostname" column gets the asset name; "IP" gets the same
        # unless we have hostname info (the build context doesn't split
        # them out — both go in asset_value).
        out.append(
            {
                "no": i,
                "hostname": asset,
                "ip": asset,
                "critical": c.get("critical", 0),
                "high": c.get("high", 0),
                "medium": c.get("medium", 0),
                "low": c.get("low", 0),
            }
        )
    out.sort(key=lambda r: r["hostname"])
    # re-number after sort
    for i, row in enumerate(out, 1):
        row["no"] = i
    return out


def _fill_summary_table(doc: DocumentObject, ctx: dict[str, Any]) -> None:
    """Fill table 4 (Summary of key findings)."""
    if len(doc.tables) < 5:
        return
    t = doc.tables[4]
    # Remove all data rows (keep header)
    for r in list(t.rows)[1:]:
        r._element.getparent().remove(r._element)
    rows = _build_host_summary(ctx)
    if not rows:
        # Always show at least one row indicating "no findings"
        new_row = t.add_row()
        _set_cell_text(new_row.cells[0], "1")
        for j in range(1, 7):
            _set_cell_text(new_row.cells[j], "0")
        return
    for r in rows:
        new_row = t.add_row()
        _set_cell_text(new_row.cells[0], str(r["no"]))
        _set_cell_text(new_row.cells[1], r["hostname"])
        _set_cell_text(new_row.cells[2], r["ip"])
        _set_cell_text(new_row.cells[3], str(r["critical"]))
        _set_cell_text(new_row.cells[4], str(r["high"]))
        _set_cell_text(new_row.cells[5], str(r["medium"]))
        _set_cell_text(new_row.cells[6], str(r["low"]))


# ---------------------------------------------------------------------------
# Replace hardcoded template literals with engagement-specific values
# ---------------------------------------------------------------------------


def _replace_template_literals(doc: DocumentObject, ctx: dict[str, Any]) -> None:
    """Replace DMC template's hardcoded literal values with the new scan's data.

    The DMC template has 'Data Management Center', 'DMC', the old test date
    ('2nd October, 2025'), and the authoring house ('Technovage Solution')
    hardcoded in prose paragraphs. The renderer swaps the tables (title
    page, scope, summary) but the paragraphs are untouched. This pass walks
    the body and replaces those literals with values from the engagement
    context.

    Must be called RIGHT AFTER ``Document(template_path)`` and BEFORE any
    other modification. Doing it at load time guarantees we only touch
    template content — analyst-written ``exec_summary`` paragraphs and the
    rebuilt detailed-finding tables are added later and stay untouched.

    Conservative by design:
      * Walks only ``doc.paragraphs`` (body) — never enters table cells,
        which are owned by ``_fill_*`` functions.
      * Skips paragraphs containing CVE references or Tenable plugin
        references (defensive — they shouldn't be in prose but if they
        are, we leave them alone).
      * Skips the disclaimer paragraph (identified by the unique phrase
        'This publication has been carefully prepared') so the brand and
        legal text are preserved.
    """
    eng = ctx.get("engagement") or {}
    client = (eng.get("client") or eng.get("name") or "").strip() or "the client"
    start_date = _format_date(eng.get("start_date") or "")
    company_name = (settings.report_company_name or "").strip() or "the testing organization"

    # Regex patterns. Word-bounded DMC prevents matching "DMC-STC-DA" etc.
    # Order doesn't matter — each pattern targets distinct text.
    replacements: list[tuple[re.Pattern[str], str]] = [
        (re.compile(r"\bData Management Center\b"), client),
        (re.compile(r"\bTechnovage Solution\b"), company_name),
        (re.compile(r"\b2nd October, 2025\b"), start_date or "the test period"),
        (re.compile(r"\bDMC\b"), client),
    ]

    def _is_disclaimer(text: str) -> bool:
        # The disclaimer is identified by a phrase unique to it. The same
        # paragraph also mentions "Technovage Solution" twice — leaving
        # that paragraph untouched preserves the brand reference.
        return "This publication has been carefully prepared" in text

    def _is_protected(text: str) -> bool:
        # Defensive: don't touch prose that looks like a finding reference.
        return (
            re.search(r"CVE-\d{4}-\d+", text) is not None
            or "Tenable Plugin" in text
            or "Plugin ID" in text
        )

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
        if _is_disclaimer(text):
            continue
        if _is_protected(text):
            continue
        new_text = text
        for pattern, replacement in replacements:
            new_text = pattern.sub(replacement, new_text)
        if new_text == text:
            continue
        # Replace the paragraph text by clearing runs and writing a new one
        # into the first run (preserves the paragraph's character style).
        for run in list(paragraph.runs):
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


# ---------------------------------------------------------------------------
# Detailed Findings: fill the template's existing tables in place
# ---------------------------------------------------------------------------
# The DMC template already ships 21 detailed-finding tables between
# the "DETAILED FINDINGS" and "CONCLUSION" headings (11 7-column tables
# with separate Impact/Recommendation columns, and 10 6-column tables
# that combine them). The renderer used to delete them all and clone
# table 5 for every chunk of 3 findings — that broke table-style
# references and created a long cascade of duplicate tables. The
# new strategy: keep every table, clear their data rows, and write
# the new findings into the existing rows. The template's column
# widths, borders, header colors, and row heights are all preserved.

SEVERITY_HEADINGS = [
    ("critical", "1.1 Critical"),
    ("high", "1.2 High"),
    ("medium", "1.3 Medium"),
    ("low", "1.4 Low"),
]
ROWS_PER_TABLE = 3  # the DMC template's detailed-finding table fits 3 data rows


def _find_detailed_finding_tables(doc: DocumentObject) -> list[Table]:
    """Return all tables between the 'DETAILED FINDINGS' heading and the
    'CONCLUSION' heading, in document order.

    Both 7-column tables (Severity / IP / Port / Issues / Impact /
    Recommendation / Action urgency) and 6-column tables (Severity /
    Hostname / Port / Issues / Recommendation and actions / Action
    urgency) live in this section. We return all of them and let the
    caller pick the right one for the current row layout.
    """
    found: list[Table] = []
    in_section = False
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if text.upper().startswith("DETAILED FINDINGS"):
                in_section = True
            elif in_section and text.upper().startswith("CONCLUSION"):
                in_section = False
                break
        elif tag == "tbl" and in_section:
            found.append(Table(child, doc))
    return found


def _clear_data_rows(table: Table) -> None:
    """Remove every row except the header from a table.

    Preserves the header row's style, the table's column widths, and
    every cell's existing ``tcPr`` properties. Only the data rows
    themselves are touched.
    """
    for r in list(table.rows)[1:]:
        r._element.getparent().remove(r._element)


def _add_detailed_row(table: Table, group: dict[str, Any]) -> None:
    """Append a single 7-column data row to a detailed-finding table.

    The table's existing header row is left intact; we add new rows
    via ``table.add_row()`` which copies the column structure. The
    cells' existing padding / borders / fonts are preserved because
    we only mutate text via ``_set_cell_text`` / ``_set_cell_multiline``.

    After the software-key merge in ``build_report_context``, a group's
    ``issues_title`` is the most-descriptive merged title (e.g.
    "OpenSSH < 9.6 Multiple Vulnerabilities") and ``cve_id`` is the
    comma-joined union of all merged CVEs. We compose the Issues cell
    from those, NOT from the unmerged ``issues_text`` (which would
    point at the wrong description and title). If the merge didn't
    happen, the pre-built ``issues_text`` is used as-is.
    """
    new_row = table.add_row()
    cells = new_row.cells
    _set_cell_text(cells[0], group["severity"].title())
    affected = "\n".join(group.get("affected_assets") or [])
    _set_cell_multiline(cells[1], affected or "N / A")
    port = group.get("port")
    protocol = group.get("protocol")
    port_str = (f"{port} / {protocol}" if protocol else str(port)) if port else "N / A"
    _set_cell_text(cells[2], port_str)
    # If this group is the result of a merge, prefer the merged
    # title + combined CVEs over the unmerged issues_text.
    if group.get("merged_count", 1) > 1:
        title = group.get("issues_title") or group.get("title") or ""
        cves = group.get("cve_id") or ""
        # Use the cleaned description from the first group in the merge
        desc = _clean_scanner_output(group.get("description") or "")
        parts = [f"{cves} {title}".strip()] if cves else [title]
        if desc:
            parts.append("")
            parts.append(desc)
        issues_text = "\n".join(parts)
    else:
        issues_text = group.get("issues_text") or group.get("issues_title") or ""
        if group.get("cve_id") and group["cve_id"] not in issues_text:
            issues_text = f"{group['cve_id']} {issues_text}".strip()
    _set_cell_multiline(cells[3], issues_text or "(no description)")
    # Impact and Recommendation come pre-populated from
    # ``build_report_context`` (the AI draft OR the suggestions-engine
    # fallback). The renderer no longer falls back to Nessus
    # boilerplate — if a row still has nothing, leave the cell empty
    # rather than dumping "See Nessus plugin output…".
    _set_cell_multiline(cells[4], group.get("impact") or "")
    _set_cell_multiline(cells[5], group.get("recommendation") or "")
    _set_cell_text(cells[6], group.get("action_urgency") or "Normal")


def _insert_heading_before_table(doc: DocumentObject, table: Table, text: str) -> None:
    """Insert a styled paragraph just before a table.

    The DMC template's body has empty paragraphs between the H1
    "DETAILED FINDINGS" heading and the first table; we drop the
    severity-section heading (``1.1 Critical`` etc.) into that gap.
    The H1 top-level heading is never recreated — it already exists
    in the template and we never touch it.
    """
    tbl_el = table._element
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p_el = p._element
    p_el.getparent().remove(p_el)
    tbl_el.addprevious(p_el)


def _rebuild_detailed_findings(doc: DocumentObject, ctx: dict[str, Any]) -> None:
    """Fill the DMC template's existing detailed-finding tables.

    Strategy:
      1. Find every 7-column table between DETAILED FINDINGS and
         CONCLUSION. These already exist in the template — we do
         NOT add or remove tables.
      2. Clear their data rows (keep the header).
      3. Group findings by severity (critical → high → medium → low).
      4. Walk through the 7-col tables, filling each with up to
         ROWS_PER_TABLE findings. When a table is full, move to the
         next. If we need more rows than the natural capacity of the
         existing tables, the last table simply gets extra rows
         (its borders / fonts / widths are preserved because we use
         ``table.add_row()``).
      5. Insert the "1.x Severity" heading before the first table
         of each severity band. The H1 "DETAILED FINDINGS" heading
         is left untouched.

    Empty severity sections are skipped (no heading inserted, table
    stays empty) — matching the DMC report's "no findings of that
    severity → no section" rule.
    """
    detail_tables = _find_detailed_finding_tables(doc)
    if not detail_tables:
        return
    # Use the 7-column tables (the rest are 6-column variants with
    # a combined "Recommendation and actions" cell — they don't match
    # our 7-cell row layout). We only use the 7-col tables and leave
    # the 6-col ones empty in place.
    seven_col_tables = [t for t in detail_tables if len(t.columns) == 7]
    if not seven_col_tables:
        return

    # Clear the data rows of EVERY table in the section, not just the
    # 7-col ones we use. The DMC template ships with sample data
    # (old engagement's hostnames, "Plugin - 12345" lines, etc.)
    # baked into the 6-col tables. Leaving that data in the report
    # would leak the original sample report's content into every
    # rendered docx. Clearing is the right call: an empty table is
    # far less embarrassing than one with someone else's findings.
    # Only the header row is preserved.
    for t in detail_tables:
        _clear_data_rows(t)

    groups_by_sev: dict[str, list[dict[str, Any]]] = {
        "critical": [],
        "high": [],
        "medium": [],
        "low": [],
    }
    for grp in ctx.get("detailed_findings", []):
        sev = (grp.get("severity") or "info").lower()
        if sev in groups_by_sev:
            groups_by_sev[sev].append(grp)

    # Track the set of finding identifiers we've already rendered so
    # we never emit the same finding twice (defensive — duplicates
    # shouldn't reach this point after the software-key merge pass,
    # but a missing dedup earlier in the pipeline shouldn't break the
    # report).
    seen_keys: set[tuple[str, str | None]] = set()

    def _key(g: dict[str, Any]) -> tuple[str, str | None]:
        return (str(g.get("vuln_id") or ""), str(g.get("port")) if g.get("port") is not None else None)

    table_iter = iter(seven_col_tables)
    current_table: Table | None = next(table_iter, None)
    rows_in_current = 0
    insertion_anchor: Table | None = None

    for sev, heading in SEVERITY_HEADINGS:
        groups = groups_by_sev.get(sev, [])
        if not groups:
            continue
        # Move to the next empty table for this severity
        if current_table is None:
            # No more tables — bail (shouldn't happen given the
            # template has 11 7-col tables).
            break
        if insertion_anchor is None:
            # First severity in the report — insert the heading
            # before the current table and use it as the first
            # target.
            _insert_heading_before_table(doc, current_table, heading)
            insertion_anchor = current_table
        else:
            # Subsequent severity. We need a fresh table for the new
            # severity's heading so the "1.x Severity" label actually
            # describes the rows that follow it. The previous
            # implementation always inserted the heading in front of
            # ``current_table`` regardless of whether the table already
            # held rows from the previous severity — that caused the
            # stacked-heading bug (two severity headings in front of the
            # same table). The fix: advance to the next empty table
            # whenever the current one already has rows.
            if rows_in_current > 0:
                # Current table is either partially or fully filled.
                # Either way, the new severity's rows should land in a
                # fresh table whose heading describes only them.
                current_table = next(table_iter, None)
                rows_in_current = 0
            if current_table is None:
                break
            _insert_heading_before_table(doc, current_table, heading)

        for grp in groups:
            k = _key(grp)
            if k in seen_keys:
                continue
            seen_keys.add(k)
            if current_table is None:
                break
            _add_detailed_row(current_table, grp)
            rows_in_current += 1
            if rows_in_current >= ROWS_PER_TABLE:
                # Move to the next table for more rows of this severity
                current_table = next(table_iter, None)
                rows_in_current = 0


# ---------------------------------------------------------------------------
# Signature footer + custom XML part (unchanged)
# ---------------------------------------------------------------------------


def _add_signature_footer_and_xml(doc: DocumentObject, signed: dict | None) -> None:
    if not signed:
        return

    section = doc.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0]
    run = p.add_run(
        f"VAPT Platform — signed by {signed.get('signer', 'unknown')} "
        f"on {signed.get('signed_at', '')} — SHA256 {signed.get('sha256', '')[:16]}…"
    )
    run.italic = True
    run.font.size = Pt(7)

    sig_json = json.dumps(signed, separators=(",", ":"))
    sig_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<vapt-signature xmlns="https://vapt.example/sig/1.0">'
        f"{sig_json}"
        "</vapt-signature>"
    ).encode()
    part_uri = PackURI("/vapt-signature.xml")
    sig_part = Part(
        part_uri,
        "application/vnd.vapt-signature+xml",
        sig_xml,
        doc.part.package,
    )
    doc.part.relate_to(sig_part, "http://vapt.example/relationships/signature")

    cp = doc.core_properties
    cp.category = "vapt-signed-report"
    note = signed.get("note") or ""
    cp.comments = note[:255]


# ---------------------------------------------------------------------------
# From-scratch fallback (used only when the template is missing)
# ---------------------------------------------------------------------------


def _render_from_scratch(
    ctx: dict[str, Any],
    *,
    signed: dict | None = None,
    exec_summary: str | None = None,
) -> bytes:
    doc = Document()
    title = doc.add_heading(f"{ctx['engagement']['name']} — Penetration Test Report", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"Client: {ctx['engagement']['client']}    Code: {ctx['engagement']['code']}\n"
    ).italic = True
    sub.add_run(
        f"Methodology: {ctx['engagement']['methodology']}    Type: {ctx['engagement']['type']}\n"
    ).italic = True
    sub.add_run(f"Generated: {ctx['generated_at']}\n").italic = True

    doc.add_heading("1. Executive Summary", level=1)
    s = ctx["summary"]
    p = doc.add_paragraph()
    p.add_run(
        f"During the engagement {s['total_findings']} findings were identified across "
        f"{s['assets_affected']} assets, representing {s['unique_vulnerabilities']} unique vulnerabilities."
    )
    overall = (s.get("overall_rating") or _overall_rating(s.get("by_severity", {})))
    p2 = doc.add_paragraph()
    p2.add_run(f"Overall security posture: {overall}.").bold = True
    if exec_summary:
        doc.add_heading("Analyst narrative", level=2)
        for para in exec_summary.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.add_paragraph().add_run("Severity breakdown:").bold = True
    for sev, n in sorted(s["by_severity"].items(), key=lambda kv: SEV_ORDER.get(kv[0], 9)):
        run = doc.add_paragraph(style="List Bullet").add_run(f"{sev.title()}: {n}")
        run.bold = True

    doc.add_heading("2. Findings", level=1)
    for i, f in enumerate(ctx["findings"], 1):
        doc.add_heading(f"{i}. [{f['severity'].upper()}] {f['title']}", level=2)
        meta = doc.add_paragraph()
        meta.add_run(f"Asset: {f['asset_value']} ({f['asset_type']})").bold = True
        if f["port"]:
            meta.add_run(f"    Port: {f['port']}/{f['protocol']}")
        if f["cve_id"]:
            meta.add_run(f"    CVE: {f['cve_id']}")
        if f["cwe_id"]:
            meta.add_run(f"    CWE: {f['cwe_id']}")
        if f["cvss_score"]:
            meta.add_run(f"    CVSS: {f['cvss_score']}")
        doc.add_heading("Description", level=3)
        doc.add_paragraph(f.get("issues_text") or f["description"] or "(no description)")
        doc.add_heading("Impact", level=3)
        doc.add_paragraph(f["impact"] or "(impact narrative pending — analyst review required)")
        doc.add_heading("Recommendation", level=3)
        doc.add_paragraph(
            f["recommendation"] or "(recommendation pending — analyst review required)"
        )

    _add_signature_footer_and_xml(doc, signed)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# public entry point
# ---------------------------------------------------------------------------


def _inject_exec_summary(doc: DocumentObject, ctx: dict[str, Any]) -> None:
    """If `ctx["exec_summary"]` is set, insert (or REPLACE) a narrative
    section immediately before the "DETAILED FINDINGS" heading.

    The replace-not-append behavior is critical: this function is called on
    every ``render_docx`` invocation (previews and signed exports alike),
    so an unconditional insert would duplicate the narrative on every
    re-render. We detect a previously-injected narrative by searching
    for the "1. Analyst Executive Narrative" heading we wrote on the
    previous pass, then remove every paragraph between that heading and
    the next section boundary before writing the new narrative.
    """
    text = ctx.get("exec_summary")
    if not isinstance(text, str) or not text.strip():
        return
    body = doc.element.body

    # Locate the DETAILED FINDINGS anchor — this is where the narrative
    # gets inserted. Also locate any pre-existing narrative heading we
    # wrote on a previous render, so we can wipe it before writing anew.
    detailed_findings_target = None
    existing_heading = None
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag != "p":
            continue
        txt = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if detailed_findings_target is None and txt.upper().startswith("DETAILED FINDINGS"):
            detailed_findings_target = child
        if existing_heading is None and "Analyst Executive Narrative" in txt:
            existing_heading = child

    if detailed_findings_target is None:
        # The template is missing the DETAILED FINDINGS heading; nothing
        # to anchor against. Bail — the rest of the renderer will still
        # populate the detailed findings tables.
        return

    # If we found a previous narrative, drop every paragraph between the
    # existing heading and the next major boundary (DETAILED FINDINGS
    # heading OR any element that's not a <w:p> — a table for example).
    # This ensures re-renders replace, not append.
    if existing_heading is not None:
        nxt = existing_heading.getnext()
        while nxt is not None:
            tag = nxt.tag.split("}")[-1]
            if tag != "p":
                break
            nxt_txt = "".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip()
            if nxt_txt.upper().startswith("DETAILED FINDINGS"):
                break
            to_remove = nxt
            nxt = nxt.getnext()
            to_remove.getparent().remove(to_remove)
        # Remove the stale heading itself.
        existing_heading.getparent().remove(existing_heading)

    # Build the new narrative: heading + one paragraph per \n\n block.
    # We construct each paragraph via the doc so character formatting is
    # applied, then detach and re-insert at the anchor position.
    insert_anchor = detailed_findings_target
    heading = doc.add_paragraph()
    hrun = heading.add_run("1. Analyst Executive Narrative")
    hrun.bold = True
    hrun.font.size = Pt(14)
    heading._element.getparent().remove(heading._element)
    insert_anchor.addprevious(heading._element)
    for para in text.split("\n\n"):
        body_text = para.strip()
        if not body_text:
            continue
        p = doc.add_paragraph(body_text)
        p._element.getparent().remove(p._element)
        insert_anchor.addprevious(p._element)


def render_docx(
    ctx: dict[str, Any],
    *,
    signed: dict | None = None,
    exec_summary: str | None = None,
) -> bytes:
    """Render the report to docx.

    Loads the DMC template from ``settings.report_template_path`` (or the
    bundled copy under ``templates/``), populates the static tables with
    engagement data, and rebuilds the Detailed Findings section. If the
    template is unavailable the function falls back to building the
    report from scratch.

    If ``signed`` is provided, embed the cryptographic signature as a
    Custom XML part (``/vapt-signature.xml``) AND a visible footer.
    """
    template_path = _resolve_template_path()
    if not template_path:
        return _render_from_scratch(ctx, signed=signed, exec_summary=exec_summary)

    doc = Document(template_path)
    _replace_template_literals(doc, ctx)
    _fill_title_page(doc, ctx)
    _fill_scope_table(doc, ctx)
    _fill_summary_table(doc, ctx)
    _inject_exec_summary(doc, ctx)
    _rebuild_detailed_findings(doc, ctx)
    _add_signature_footer_and_xml(doc, signed)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HTML preview
# ---------------------------------------------------------------------------
#
# Single source of truth: the preview is generated by rendering the docx
# (the same function the signed export uses) and converting it to HTML via
# mammoth. This guarantees the preview and the rendered docx can never
# drift apart — what the analyst sees in the browser is exactly what will
# come out of POST /reports/{rid}/render.
#
# The preview never embeds the cryptographic signature: the report is not
# considered signed until it has gone through the render-and-approve
# pipeline. We always call render_docx with signed=None.


log = structlog.get_logger(__name__)


_SEVERITY_LABEL_TO_CLASS = {
    "Critical": "sev-critical",
    "High": "sev-high",
    "Medium": "sev-medium",
    "Low": "sev-low",
    "Informational": "sev-info",
}


def _add_severity_classes(html: str) -> str:
    """Tag severity cells with CSS classes for visual color-coding.

    The docx renderer writes the severity label as the bare text content
    of a single-cell table row (e.g. ``<td><p>Critical</p></td>``). We
    post-process those exact cells to add a class so the wrapper CSS can
    color them. Non-severity cells are untouched.
    """
    for label, css_class in _SEVERITY_LABEL_TO_CLASS.items():
        pattern = r"<td><p>" + re.escape(label) + r"</p></td>"
        replacement = f'<td class="{css_class}"><p>{label}</p></td>'
        html = re.sub(pattern, replacement, html)
    return html


def _tag_editable_regions(body_html: str, ctx: dict[str, Any]) -> str:
    """Add ``data-field`` attributes to specific elements so the editor
    can map DOM mutations back to the structured draft fields.

    Tags applied:
      * the analyst-narrative paragraphs (between the
        "1. Analyst Executive Narrative" heading and the next ``<h1>``)
        are wrapped in a single ``<div data-field="exec_summary">`` so
        the whole block is one editable unit that preserves paragraph
        breaks.
      * the 5th ``<td>`` (impact) and 6th ``<td>`` (recommendation) of
        every data row in the Detailed Findings tables are tagged with
        ``data-field="finding.impact"`` / ``"finding.recommendation"``
        plus ``data-finding-id="<finding_id>"`` (the key the editor's
        override map is indexed by). The detailed-findings rows are
        emitted by ``_rebuild_detailed_findings`` in the same order as
        ``ctx["detailed_findings"]``, so we pair them by index. When a
        group collapses multiple findings that share ``(vuln_id, port)``
        the first finding wins — this matches what ``_apply_draft``
        surfaces in the rendered docx anyway.

    The editor only enables ``contenteditable`` on these elements when
    the user opts in via the "Edit in place" toggle; the HTML payload
    itself stays read-only.

    ``body_html`` is the inner body content produced by mammoth (no
    surrounding ``<html>``/``<body>`` tags). On parse failure the input
    is returned unchanged so a malformed mammoth output never breaks
    preview rendering.
    """
    if not body_html:
        return body_html

    try:
        wrapped = f"<div id='vapt-edit-root'>{body_html}</div>"
        doc = lxml_html.fragment_fromstring(wrapped, create_parent=True)
    except Exception:
        return body_html

    # 1. Wrap the executive narrative paragraphs in a single <div data-field="exec_summary">.
    for p in doc.iter("p"):
        if p.text_content().strip() == "1. Analyst Executive Narrative":
            parent = p.getparent()
            if parent is None:
                break
            siblings = list(parent)
            idx = siblings.index(p)
            to_wrap = []
            j = idx + 1
            while j < len(siblings) and siblings[j].tag == "p":
                to_wrap.append(siblings[j])
                j += 1
            if to_wrap:
                wrapper = lxml_html.Element("div")
                wrapper.set("data-field", "exec_summary")
                for child in to_wrap:
                    parent.remove(child)
                    wrapper.append(child)
                parent.insert(idx + 1, wrapper)
            break

    # 2. Tag finding impact / recommendation cells.
    # Walk the document, tracking the section between DETAILED FINDINGS and CONCLUSION.
    detailed_findings = ctx.get("detailed_findings") or []
    if detailed_findings:
        # Build a (vuln_id, port) -> first finding_id map. The editor's
        # override map is keyed by finding_id, but the detailed-findings
        # groups collapse multiple findings that share (vuln_id, port);
        # the WYSIWYG edit targets the first finding in that group, which
        # is what `_apply_draft` will surface in the rendered docx anyway.
        group_key_to_fid: dict[tuple[str, str], str] = {}
        for row in ctx.get("findings") or []:
            vid = str(row.get("vuln_id") or "")
            fid = str(row.get("finding_id") or "")
            if not vid or not fid:
                continue
            port = row.get("port")
            port_key = "" if port is None else str(port)
            group_key_to_fid.setdefault((vid, port_key), fid)

        in_section = False
        data_rows: list[Any] = []
        for elem in doc.iter():
            tag = elem.tag
            if tag == "h1":
                txt = elem.text_content().strip()
                if "DETAILED FINDINGS" in txt:
                    in_section = True
                elif "CONCLUSION" in txt and in_section:
                    in_section = False
                    break
            elif tag == "table" and in_section:
                for i, tr in enumerate(elem.iter("tr")):
                    if i == 0:
                        continue  # header row
                    data_rows.append(tr)
                    if len(data_rows) >= len(detailed_findings):
                        break
                if len(data_rows) >= len(detailed_findings):
                    break

        for tr, finding in zip(data_rows, detailed_findings):
            vid = str(finding.get("vuln_id") or "")
            port = finding.get("port")
            port_key = "" if port is None else str(port)
            fid = group_key_to_fid.get((vid, port_key), "")
            if not fid:
                continue
            cells = list(tr.iter("td"))
            if len(cells) >= 6:
                cells[4].set("data-field", "finding.impact")
                cells[4].set("data-finding-id", fid)
                cells[5].set("data-field", "finding.recommendation")
                cells[5].set("data-finding-id", fid)

    new_body_bytes = lxml_html.tostring(doc, encoding="utf-8")
    new_body = new_body_bytes.decode("utf-8") if isinstance(new_body_bytes, bytes) else new_body_bytes
    # Strip the synthetic wrapper we added for parsing.
    new_body = re.sub(
        r'^<div id="vapt-edit-root">(.*)</div>$',
        r"\1",
        new_body,
        count=1,
        flags=re.DOTALL,
    )
    if new_body == body_html:
        return body_html
    return new_body


_PREVIEW_CSS = """
/* A4 page sizing (8.27in × 11.16in) with 1in top/bottom, 0.75in left/right
   margins — mirrors the DMC template's section properties, which mammoth
   strips during conversion. box-sizing: border-box keeps the body within
   the 8.27in A4 width inclusive of padding. */
body {
  font-family: 'Calibri', 'Helvetica', 'Arial', sans-serif;
  font-size: 11pt;
  line-height: 1.45;
  color: #1a1a1a;
  max-width: 8.27in;
  margin: 0 auto;
  padding: 1in 0.75in;
  background: #fff;
  box-sizing: border-box;
}
h1 {
  color: #1F3864 !important;
  font-size: 18pt !important;
  font-weight: 700 !important;
  border-bottom: 2px solid #1F3864 !important;
  margin: 0.4in 0 0.15in !important;
  padding-bottom: 0.05in !important;
  text-transform: uppercase !important;
  letter-spacing: 0.02em !important;
}
h2 {
  color: #1F3864;
  font-size: 14pt;
  margin: 0.3in 0 0.1in;
  padding-bottom: 0.03in;
  border-bottom: 1px solid #cbd5e1;
}
h3 { font-size: 12pt; margin: 0.2in 0 0.08in; }
p { margin: 0.1in 0; }
table {
  width: 100%;
  border-collapse: collapse;
  margin: 0.2in 0;
  font-size: 10pt;
  table-layout: auto;
  border: 2px solid #1F3864;
}
th {
  background: #1F3864;
  color: #fff;
  font-weight: 700;
  border: 1px solid #1F3864;
  padding: 8pt 10pt;
  text-align: left;
  font-size: 10pt;
  letter-spacing: 0.02em;
}
td {
  border: 1px solid #c0c0c0;
  padding: 8pt 10pt;
  vertical-align: top;
  word-wrap: break-word;
  white-space: pre-wrap;
  font-size: 10pt;
  line-height: 1.4;
}
td p { margin: 0 0 6pt 0; }
td p:last-child { margin-bottom: 0; }
tr:nth-child(even) td { background: #f3f6fa; }
tr:nth-child(odd) td { background: #ffffff; }

/* Detailed-findings table layout: 7 columns. The Issues column has the
   long content (CVE + multi-line plugin description) so give it more
   room; the Severity / Action urgency cells are short. */
table:not(:first-of-type) tr td:nth-child(1),
table:not(:first-of-type) tr th:nth-child(1) { width: 9%; }   /* Severity */
table:not(:first-of-type) tr td:nth-child(2),
table:not(:first-of-type) tr th:nth-child(2) { width: 12%; }  /* Affected IP/URL */
table:not(:first-of-type) tr td:nth-child(3),
table:not(:first-of-type) tr th:nth-child(3) { width: 8%; }   /* Affected Port */
table:not(:first-of-type) tr td:nth-child(4),
table:not(:first-of-type) tr th:nth-child(4) { width: 30%; }  /* Issues */
table:not(:first-of-type) tr td:nth-child(5),
table:not(:first-of-type) tr th:nth-child(5) { width: 15%; }  /* Impact */
table:not(:first-of-type) tr td:nth-child(6),
table:not(:first-of-type) tr th:nth-child(6) { width: 22%; }  /* Recommendation */
table:not(:first-of-type) tr td:nth-child(7),
table:not(:first-of-type) tr th:nth-child(7) { width: 4%; }   /* Action urgency */
/* Cover page treatment: the first <table> in the docx is the title page.
   The DMC template's cover carries the Technovage logo, the client
   name, and a subtle blue gradient — match that here. The left cell
   holds the "For/Type/From/Date" labels (white-on-blue monospace);
   the right cell holds the engagement info. */
table:first-of-type {
  border: 2px solid #1F3864;
  margin-bottom: 0.3in;
  background: linear-gradient(180deg, #ffffff 0%, #f0f4fa 100%);
  border-collapse: separate;
  border-spacing: 0;
}
table:first-of-type td {
  background: transparent;
  border: none;
  padding: 14pt 12pt;
  vertical-align: top;
}
table:first-of-type td:first-child {
  background: #1F3864;
  color: #fff;
  font-family: 'Courier New', 'Consolas', monospace;
  font-weight: 700;
  white-space: pre;
  width: 32%;
}
table:first-of-type td:last-child {
  font-size: 13pt;
  line-height: 1.7;
}
img { max-width: 100%; height: auto; }
.sev-critical { background: #dc2626; color: #fff; font-weight: 700; padding: 4pt 8pt; border-radius: 3pt; display: inline-block; }
.sev-high     { background: #ea580c; color: #fff; font-weight: 700; padding: 4pt 8pt; border-radius: 3pt; display: inline-block; }
.sev-medium   { background: #d97706; color: #fff; font-weight: 700; padding: 4pt 8pt; border-radius: 3pt; display: inline-block; }
.sev-low      { background: #ca8a04; color: #fff; font-weight: 700; padding: 4pt 8pt; border-radius: 3pt; display: inline-block; }
.sev-info     { background: #2563eb; color: #fff; font-weight: 700; padding: 4pt 8pt; border-radius: 3pt; display: inline-block; }

/* The Issues cell has long multi-line content (CVE + plugin description).
   Use a monospace font and tighter line-height for readability. */
table:not(:first-of-type) tr td:nth-child(4) {
  font-family: 'Consolas', 'Menlo', 'Courier New', monospace;
  font-size: 9pt;
  line-height: 1.35;
  white-space: pre-wrap;
}
"""


def render_preview_html(ctx: dict[str, Any]) -> str:
    """Return a self-contained HTML preview of the report.

    The preview is a faithful HTML rendering of the docx that
    ``POST /reports/{rid}/render`` would produce: the docx IS the source
    of truth, the preview is just its HTML projection. Because the two
    paths share the same renderer, the preview and the final docx can
    never disagree on content.

    The preview is read-only (it does not burn a version number) and
    never embeds the cryptographic signature. The signed footer /
    custom XML part is only added by ``render_docx(..., signed=...)``
    during the render-and-approve flow.
    """
    docx_bytes = render_docx(ctx, signed=None)

    # mammoth's default image handler is ``mammoth.images.data_uri``, which
    # already base64-encodes every embedded image and emits an <img src=
    # "data:image/...;base64,...">. We rely on that default rather than
    # passing our own convert_image callback: mammoth 1.x's
    # ``images.img_element`` decorator accesses ``image.alt_text`` on its
    # argument, and a misplaced second positional arg to ``convert_to_html``
    # routes the wrapped function through ``transform_document`` — which
    # passes a Document, not an Image, and crashes. Using the default
    # produces all 8 template images as data URIs, which the sandboxed
    # iframe can render without external requests.
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    body = result.value
    for msg in result.messages:
        log.warning(
            "mammoth_preview_warning",
            type=msg.type,
            message=msg.message,
        )

    body = _add_severity_classes(body)
    body = _tag_editable_regions(body, ctx)

    return (
        "<!DOCTYPE html>\n"
        '<html lang="en">\n'
        "<head>\n"
        '  <meta charset="utf-8">\n'
        "  <style>\n"
        f"{_PREVIEW_CSS}\n"
        "  </style>\n"
        "</head>\n"
        f"<body>\n{body}\n</body>\n"
        "</html>\n"
    )
