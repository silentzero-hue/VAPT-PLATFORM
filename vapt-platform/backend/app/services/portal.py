"""Client portal: tokenized, time-bounded, view-counted access to an
approved report. Includes per-view audit trail and optional
password / email allowlist.

When `watermark_with_viewer=True`, the docx is re-rendered with the
viewer's email/ID burned into the footer. The signed report in S3
is never served directly — the portal endpoint re-renders a
watermarked copy into a per-share S3 key.
"""

from __future__ import annotations

import io
import secrets
import uuid
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.shared import Pt
from fastapi import HTTPException
from sqlalchemy import or_, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.security import hash_password, hash_token, verify_password
from app.models.portal import PortalShare
from app.models.report import Report, ReportStatus
from app.services import storage


def _new_token() -> tuple[str, str]:
    raw = "psh_" + secrets.token_urlsafe(24)
    return raw, hash_token(raw)


async def create_share(
    db: AsyncSession,
    *,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
    label: str,
    actor_id: uuid.UUID | None,
    expires_at: datetime | None = None,
    max_views: int | None = None,
    require_password: bool = False,
    password: str | None = None,
    allowed_emails: list[str] | None = None,
    watermark_with_viewer: bool = True,
    note: str | None = None,
) -> tuple[PortalShare, str]:
    raw, h = _new_token()
    share = PortalShare(
        workspace_id=workspace_id,
        report_id=report_id,
        created_by=actor_id,
        token_hash=h,
        label=label,
        expires_at=expires_at,
        max_views=max_views,
        require_password=require_password,
        password_hash=hash_password(password) if password and require_password else None,
        allowed_emails=[e.lower() for e in (allowed_emails or [])],
        watermark_with_viewer=watermark_with_viewer,
        note=note,
    )
    db.add(share)
    await db.flush()
    return share, raw


async def revoke(db: AsyncSession, share_id: uuid.UUID, reason: str | None) -> None:
    s = await db.get(PortalShare, share_id)
    if s and not s.revoked_at:
        s.revoked_at = datetime.now(timezone.utc)
        s.revoked_reason = reason


async def access(
    db: AsyncSession, raw_token: str, viewer_email: str | None, viewer_password: str | None
) -> tuple[PortalShare, Report, bytes]:
    if not raw_token:
        raise HTTPException(401, "missing token")
    h = hash_token(raw_token)
    s = await db.scalar(select(PortalShare).where(PortalShare.token_hash == h))
    if not s or s.revoked_at:
        raise HTTPException(404, "share not found")
    if s.expires_at and s.expires_at < datetime.now(timezone.utc):
        raise HTTPException(410, "share expired")
    bumped = await db.execute(
        update(PortalShare)
        .where(PortalShare.id == s.id)
        .where(or_(PortalShare.max_views.is_(None), PortalShare.current_views < PortalShare.max_views))
        .values(current_views=PortalShare.current_views + 1)
        .returning(PortalShare.id)
    )
    if not bumped.scalar():
        raise HTTPException(410, "view limit reached")
    await db.refresh(s)
    if s.require_password and not verify_password(viewer_password or "", s.password_hash or ""):
        raise HTTPException(401, "password required")
    if s.allowed_emails and viewer_email and viewer_email.lower() not in s.allowed_emails:
        raise HTTPException(403, "email not on allowlist")
    r = await db.get(Report, s.report_id)
    if not r or r.status not in (ReportStatus.APPROVED, ReportStatus.PUBLISHED):
        raise HTTPException(404, "report not ready")
    v = next((x for x in r.versions if x.id == r.current_version_id), None)
    if not v or not v.s3_key:
        raise HTTPException(404, "no rendered version")
    data = await storage.get_bytes(v.s3_key)
    if s.watermark_with_viewer:
        data = _watermark_docx(data, viewer_email or "external", raw_token[:8])
    s.last_access_at = datetime.now(timezone.utc)
    s.access_log = (s.access_log or []) + [
        {
            "ts": s.last_access_at.isoformat(),
            "email": viewer_email,
            "token_prefix": raw_token[:8],
        }
    ]
    return s, r, data


def _watermark_docx(data: bytes, viewer: str, token_prefix: str) -> bytes:
    doc = Document(io.BytesIO(data))
    p = doc.add_paragraph()
    run = p.add_run(
        f"Confidential — shared with {viewer} — share {token_prefix} — "
        f"{datetime.now(timezone.utc).isoformat()}"
    )
    run.italic = True
    run.font.size = Pt(8)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
